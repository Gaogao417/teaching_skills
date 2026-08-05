#!/usr/bin/env python3
"""Batch-transcribe all 一模/二模 docx papers under documents/初三/.

For each exam directory, pick ONE source containing BOTH questions and answers:
  - single docx (题答合一) → use it directly
  - exam + answers split (source-exam.docx + source-answers.docx) → merge into one
  - source + 教师版/解析 (教师版 has Q+A) → prefer 教师版
  - exam-only / image-only → skip
  - legacy .doc → feed to run_live_paper as source-kind "doc" (LibreOffice path)

paper_id = "<届>-<区拼音>-<YIMO|ERMO>", e.g. 2023-HUANGPU-YIMO.

Runs run_live_paper (one subprocess per paper) in an 8-worker pool. Skips papers
already transcribed (a completed transcription.yaml exists). Passes langfuse keys
when LF_PUBLIC_KEY/LF_SECRET_KEY env vars are set.

Usage:
    source ~/.zshrc 2>/dev/null
    ./.venv/bin/python -m scripts.question_transcription.workflow.cli.batch_transcribe_papers \\
        --documents-root documents/初三 \\
        --workers 8 [--limit N] [--dry-run]
"""

from __future__ import annotations

import argparse
import os
import re
import subprocess
import sys
import tempfile
from concurrent.futures import (
    ProcessPoolExecutor,
    as_completed,
)
from dataclasses import dataclass
from pathlib import Path

# --------------------------------------------------------------------------- #
# District-name pinyin (目录名 → paper_id 拼音段). 单区多音/简写统一。
# --------------------------------------------------------------------------- #
DISTRICT_PINYIN: dict[str, str] = {
    "嘉定区": "JIADING", "奉贤区": "FENGXIAN", "宝山区": "BAOSHAN",
    "宝山区嘉定区": "BAOSHAN-JIADING",  # 合考
    "崇明区": "CHONGMING", "徐汇区": "XUHUI", "普陀区": "PUTUO",
    "杨浦区": "YANGPU", "松江区": "SONGJIANG", "浦东新区": "PUDONG",
    "浦东区": "PUDONG",  # 2025届 简写
    "虹口区": "HONGKOU", "金山区": "JINSHAN", "长宁区": "CHANGNING",
    "闵行区": "MINHANG", "闸北区": "ZHABEI", "青浦区": "QINGPU",
    "静安区": "JINGAN", "黄浦区": "HUANGPU",
}


@dataclass
class PaperJob:
    paper_id: str
    source: str           # absolute path to the source file to transcribe
    source_kind: str      # "docx" | "doc"
    label: str            # directory name, for logging
    note: str = ""        # why this source was chosen / any caveat


def _year_from_dir(dirname: str) -> int | None:
    m = re.match(r"(\d{4})届", dirname)
    return int(m.group(1)) if m else None


def _mock_from_dir(dirname: str) -> str | None:
    if "一模" in dirname:
        return "YIMO"
    if "二模" in dirname:
        return "ERMO"
    return None


def _district_from_dir(dirname: str) -> str | None:
    # "<届>-上海市<区>-初三..." → 取 "上海市" 之后、"区" 之前
    m = re.search(r"上海市(.+?区)-初三", dirname)
    if not m:
        return None
    district = m.group(1)
    return DISTRICT_PINYIN.get(district)


def make_paper_id(dirname: str) -> str | None:
    year = _year_from_dir(dirname)
    mock = _mock_from_dir(dirname)
    district = _district_from_dir(dirname)
    if year and mock and district:
        return f"{year}-{district}-{mock}"
    return None


# --------------------------------------------------------------------------- #
# Source selection per directory
# --------------------------------------------------------------------------- #

Q_MARKERS = ("选择题", "填空题", "解答题")
A_MARKERS = ("参考答案", "【解析】", "【详解】", "试题解析", "试题答案")


def _docx_text(path: Path) -> str:
    """Cheap text extraction from a docx (unzip document.xml, strip tags)."""
    import zipfile

    try:
        with zipfile.ZipFile(path) as z:
            with z.open("word/document.xml") as f:
                xml = f.read().decode("utf-8", errors="ignore")
    except Exception:
        return ""
    # collapse tags to a space, decode a couple of common entities
    text = re.sub(r"<[^>]+>", " ", xml)
    text = text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    return re.sub(r"\s+", " ", text)


def _has_q_and_a(path: Path) -> tuple[bool, bool]:
    """Return (has_questions, has_answers) for a docx by marker grep."""
    text = _docx_text(path)
    has_q = any(m in text for m in Q_MARKERS)
    has_a = any(m in text for m in A_MARKERS)
    return has_q, has_a


def merge_docx(parts: list[Path], out: Path) -> Path:
    """Merge several docx into one by concatenating paragraphs (python-docx)."""
    from docx import Document

    if len(parts) == 1:
        return parts[0]
    merged = Document(str(parts[0]))
    for part in parts[1:]:
        sub = Document(str(part))
        # page break then append each paragraph
        merged.add_page_break()
        for para in sub.paragraphs:
            new = merged.add_paragraph()
            for run in para.runs:
                new.add_run(run.text)
    out.parent.mkdir(parents=True, exist_ok=True)
    merged.save(str(out))
    return out


def select_source(directory: Path, merge_dir: Path) -> PaperJob | None:
    """Pick the transcription source for one exam directory, or None to skip."""
    dirname = directory.name
    paper_id = make_paper_id(dirname)
    if not paper_id:
        return None

    # Gather candidate docx/doc files at top level (skip word/, word_*/ asset dirs).
    docx_files = sorted(p.name for p in directory.glob("*.docx"))
    doc_files = sorted(p.name for p in directory.glob("*.doc"))  # not .docx

    # --- docx present: prefer a combined Q+A one ---
    if docx_files:
        scored: list[tuple[Path, bool, bool, str]] = []
        for name in docx_files:
            p = directory / name
            has_q, has_a = _has_q_and_a(p)
            scored.append((p, has_q, has_a, name))
        # 1) a single file with BOTH Q and A
        combined = [(p, n) for (p, q, a, n) in scored if q and a]
        if len(combined) == 1:
            p, n = combined[0]
            return PaperJob(paper_id, str(p), "docx", dirname,
                            note=f"single Q+A: {n}")
        if len(combined) > 1:
            # prefer 教师版/answers/source.docx; else first
            preference_order = ("教师版", "answers", "source.docx", "精品解析")
            chosen = combined[0]
            for pref in preference_order:
                hit = [c for c in combined if pref in c[1]]
                if hit:
                    chosen = hit[0]
                    break
            p, n = chosen
            return PaperJob(paper_id, str(p), "docx", dirname,
                            note=f"Q+A pick: {n}")

        # 2) split: exam (Q only) + answers — merge them
        q_only = [(p, n) for (p, q, a, n) in scored if q and not a]
        a_any = [(p, n) for (p, q, a, n) in scored if a]
        if q_only and a_any:
            exam_p = q_only[0][0]
            ans_p = a_any[0][0]
            out = merge_dir / f"{paper_id}-merged.docx"
            merge_docx([exam_p, ans_p], out)
            return PaperJob(paper_id, str(out), "docx", dirname,
                            note=f"merged: {exam_p.name} + {ans_p.name}")

        # 3) only exam-only docx (Q, no A anywhere) — skip (no answers to transcribe)
        if scored and not any(a for (_, _, a, _) in scored):
            return None
        # 4) fallback: first docx if it has anything
        if scored:
            p, q, a, n = scored[0]
            return PaperJob(paper_id, str(p), "docx", dirname,
                            note=f"fallback: {n} (q={q},a={a})")
        return None

    # --- only legacy .doc: feed to run_live_paper as source-kind "doc" ---
    if doc_files:
        # prefer source.doc
        src = next((directory / n for n in doc_files if n == "source.doc"),
                   directory / doc_files[0])
        return PaperJob(paper_id, str(src), "doc", dirname,
                        note=f"legacy .doc: {src.name}")

    # --- image-only: skip ---
    return None


# --------------------------------------------------------------------------- #
# Already-transcribed check + worker
# --------------------------------------------------------------------------- #

REPO_ROOT = Path(__file__).resolve().parents[4]


def _transcribed(paper_id: str) -> bool:
    """True if any completed transcription.yaml exists under this paper's runs."""
    base = REPO_ROOT / "build" / "question-ingestion" / paper_id
    if not base.is_dir():
        return False
    for run_dir in base.iterdir():
        t = run_dir / "structured" / "transcription.yaml"
        if t.exists() and t.stat().st_size > 0:
            return True
    return False


def run_one(job: PaperJob, venv_python: str, lf_pub: str | None, lf_sec: str | None,
            langfuse_host: str) -> tuple[str, str]:
    """Run run_live_paper in a subprocess for one job. Returns (paper_id, status)."""
    cmd = [
        venv_python, "-m", "scripts.question_transcription.workflow.run_live_paper",
        "--paper-id", job.paper_id,
        "--source", job.source,
        "--source-kind", job.source_kind,
        "--agent-host", "claude-code",
    ]
    if lf_pub and lf_sec:
        cmd += ["--langfuse-host", langfuse_host,
                "--langfuse-public-key", lf_pub,
                "--langfuse-secret-key", lf_sec]
    log_file = REPO_ROOT / "build" / "question-ingestion" / job.paper_id / "batch.log"
    log_file.parent.mkdir(parents=True, exist_ok=True)
    try:
        with open(log_file, "w") as lf:
            proc = subprocess.run(
                cmd, cwd=str(REPO_ROOT), stdout=lf, stderr=subprocess.STDOUT,
                timeout=1800, check=False,
            )
        # transcribe_whole_paper writes structured/transcription.yaml on success
        ok = _transcribed(job.paper_id)
        status = "OK" if ok else f"FAIL(rc={proc.returncode})"
    except subprocess.TimeoutExpired:
        status = "TIMEOUT"
    except Exception as exc:  # pragma: no cover
        status = f"ERR:{type(exc).__name__}:{exc}"
    return job.paper_id, status


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #


def main(argv: list[str] | None = None) -> int:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--documents-root", default="documents/初三")
    p.add_argument("--workers", type=int, default=8)
    p.add_argument("--limit", type=int, default=0, help="process only N papers (0=all)")
    p.add_argument("--dry-run", action="store_true",
                   help="print the discovered jobs and exit (no transcription)")
    p.add_argument("--langfuse-host", default="http://localhost:3000")
    args = p.parse_args(argv)

    docs_root = (REPO_ROOT / args.documents_root).resolve()
    if not docs_root.is_dir():
        print(f"ERROR: documents root not found: {docs_root}", file=sys.stderr)
        return 2

    merge_dir = REPO_ROOT / "build" / "_merged-sources"
    venv_python = str(REPO_ROOT / ".venv" / "bin" / "python")

    # Discover
    jobs: list[PaperJob] = []
    skipped: list[str] = []
    for entry in sorted(docs_root.iterdir()):
        if not entry.is_dir():
            continue
        if "一模" not in entry.name and "二模" not in entry.name:
            continue
        job = select_source(entry, merge_dir)
        if job is None:
            skipped.append(entry.name)
            continue
        if _transcribed(job.paper_id):
            skipped.append(f"{entry.name}  [already transcribed: {job.paper_id}]")
            continue
        jobs.append(job)

    print(f"Discovered {len(jobs)} papers to transcribe; skipped {len(skipped)}.")
    if skipped:
        print("Skipped (first 30):")
        for s in skipped[:30]:
            print(f"  - {s}")
        if len(skipped) > 30:
            print(f"  ... ({len(skipped) - 30} more)")

    if args.dry_run:
        print("\nJobs:")
        for j in jobs:
            print(f"  {j.paper_id}  [{j.source_kind}]  {j.note}")
            print(f"      src: {j.source}")
        return 0

    if not jobs:
        print("Nothing to do.")
        return 0

    if args.limit:
        jobs = jobs[: args.limit]
        print(f"Limited to {len(jobs)} jobs.")

    lf_pub = os.environ.get("LF_PUBLIC_KEY") or os.environ.get("LANGFUSE_PUBLIC_KEY")
    lf_sec = os.environ.get("LF_SECRET_KEY") or os.environ.get("LANGFUSE_SECRET_KEY")

    done = 0
    ok = 0
    with ProcessPoolExecutor(max_workers=args.workers) as pool:
        futures = {
            pool.submit(run_one, j, venv_python, lf_pub, lf_sec, args.langfuse_host): j
            for j in jobs
        }
        for fut in as_completed(futures):
            j = futures[fut]
            try:
                paper_id, status = fut.result()
            except Exception as exc:  # pragma: no cover
                paper_id, status = j.paper_id, f"POOL-ERR:{exc}"
            done += 1
            if status == "OK":
                ok += 1
            print(f"[{done}/{len(jobs)}] {paper_id} -> {status}  ({j.label})")

    print(f"\nDone: {ok}/{len(jobs)} OK, {len(jobs) - ok} failed.")
    return 0 if ok == len(jobs) else 1


if __name__ == "__main__":
    raise SystemExit(main())
